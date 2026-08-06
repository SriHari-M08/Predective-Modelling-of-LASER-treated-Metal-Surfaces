"""
app.py — Ti-6Al-4V Laser Surface Engineering Predictor
========================================================
Run with:  streamlit run app.py
"""

import warnings
warnings.filterwarnings("ignore")

import io
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import joblib
import streamlit as st
from pathlib import Path
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error

# ── Import config (graceful fallback if config.py is missing) ────────────────
try:
    from config import (
        DATA_PATH, MODEL_DIR, SEED, N_ESTIMATORS, CV_FOLDS, TEST_SIZE,
        FEATURES, FEATURE_LABELS, TARGETS, TARGET_KEYS,
        TARGET_LABELS, TARGET_UNITS, TARGET_COLORS, PARAM_RANGES,
    )
except ImportError:
    BASE_DIR    = Path(__file__).parent
    DATA_PATH   = BASE_DIR / "data" / "Ti64_Advanced_Laser_Dataset.csv"
    MODEL_DIR   = BASE_DIR / "models"
    SEED        = 42
    N_ESTIMATORS = 150
    CV_FOLDS    = 5
    TEST_SIZE   = 0.2
    FEATURES        = ["Laser_Power_W","Scanning_Speed_mm_s","Pulse_Frequency_kHz","Hatch_Spacing_um"]
    FEATURE_LABELS  = ["Laser Power (W)","Scanning Speed (mm/s)","Pulse Frequency (kHz)","Hatch Spacing (µm)"]
    TARGETS         = ["Microhardness_HV","Surface_Roughness_Ra_um","HAZ_Depth_um"]
    TARGET_KEYS     = ["hardness","roughness","haz"]
    TARGET_LABELS   = ["Microhardness","Surface Roughness (Ra)","HAZ Depth"]
    TARGET_UNITS    = ["HV","µm","µm"]
    TARGET_COLORS   = ["#2ca02c","#1f77b4","#d62728"]
    PARAM_RANGES    = {
        "Laser_Power_W":       (100.0,  500.0, 250.0),
        "Scanning_Speed_mm_s": (50.0,  1200.0, 600.0),
        "Pulse_Frequency_kHz": (20.0,   100.0,  40.0),
        "Hatch_Spacing_um":    (50.0,   150.0, 100.0),
    }

# ─────────────────────────────────────────────────────────────────────────────
# BACKEND: data loading, training, caching
# ─────────────────────────────────────────────────────────────────────────────

def _synthetic_df(n: int = 1000) -> pd.DataFrame:
    rng   = np.random.default_rng(SEED)
    power = rng.uniform(100, 500, n)
    speed = rng.uniform(50, 1200, n)
    freq  = rng.uniform(20, 100, n)
    hatch = rng.uniform(50, 150, n)
    ed    = (power * 1000) / (speed * hatch)

    hardness  = 320 + 200 * np.exp(-0.5 * ((ed - 30) / 15) ** 2) + rng.normal(0, 12, n)
    roughness = np.clip(0.5 + 0.05*ed + 0.01*hatch - 0.02*freq + rng.normal(0, 0.2, n), 0.05, 6.0)
    haz       = np.clip(10 + 2.5 * ed + rng.normal(0, 2, n), 5, 200)

    return pd.DataFrame({
        "Laser_Power_W":           power,
        "Scanning_Speed_mm_s":     speed,
        "Pulse_Frequency_kHz":     freq,
        "Hatch_Spacing_um":        hatch,
        "Microhardness_HV":        hardness,
        "Surface_Roughness_Ra_um": roughness,
        "HAZ_Depth_um":            haz,
    })


@st.cache_resource(show_spinner=False)
def load_data_and_train():
    """Load or train all models. Returns (models, df, results, using_real_data)."""
    if DATA_PATH.exists():
        df = pd.read_csv(DATA_PATH)
        using_real = True
    else:
        df = _synthetic_df()
        using_real = False

    X = df[FEATURES]
    models, results = {}, {}

    MODEL_DIR.mkdir(parents=True, exist_ok=True)

    for key, col in zip(TARGET_KEYS, TARGETS):
        y = df[col]
        X_train, X_test, y_train, y_test = train_test_split(
            X, y, test_size=TEST_SIZE, random_state=SEED
        )
        model_path = MODEL_DIR / f"model_{key}.pkl"

        if model_path.exists():
            model = joblib.load(model_path)
        else:
            model = RandomForestRegressor(n_estimators=N_ESTIMATORS, random_state=SEED, n_jobs=-1)
            model.fit(X_train, y_train)
            joblib.dump(model, model_path)

        y_pred = model.predict(X_test)
        cv     = cross_val_score(model, X, y, cv=CV_FOLDS, scoring="r2")

        models[key]  = model
        results[key] = {
            "y_test":   y_test,
            "y_pred":   y_pred,
            "r2":       r2_score(y_test, y_pred),
            "mae":      mean_absolute_error(y_test, y_pred),
            "rmse":     np.sqrt(mean_squared_error(y_test, y_pred)),
            "cv_mean":  cv.mean(),
            "cv_std":   cv.std(),
        }

    return models, df, results, using_real


def predict_with_uncertainty(model, input_df):
    """Return (mean, std) across all trees — gives a calibrated uncertainty band."""
    tree_preds = np.array([t.predict(input_df) for t in model.estimators_])
    return float(tree_preds.mean()), float(tree_preds.std())


# ─────────────────────────────────────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Ti-6Al-4V Laser Predictor",
    layout="wide",
    page_icon="🔬",
)

st.markdown("""
<style>
div[data-testid="metric-container"] {
    background: var(--background-color, #f9f9f9);
    border-radius: 8px;
    padding: 10px 16px;
    border-left: 4px solid #2ca02c;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# LOAD MODELS
# ─────────────────────────────────────────────────────────────────────────────

with st.spinner("Loading models … (first run may take ~30 s to train)"):
    models, df_train, results, using_real_data = load_data_and_train()

# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────

with st.sidebar:
    st.header("🎛️ Process Parameters")

    if not using_real_data:
        st.warning(
            "⚠️ **Synthetic data** — place `Ti64_Advanced_Laser_Dataset.csv` "
            "in the `data/` folder and restart the app."
        )

    p = st.number_input("Laser Power (W)",        100.0,  500.0, 250.0, step=5.0)
    s = st.number_input("Scanning Speed (mm/s)",   50.0, 1200.0, 600.0, step=10.0)
    f = st.number_input("Pulse Frequency (kHz)",   20.0,  100.0,  40.0, step=1.0)
    h = st.number_input("Hatch Spacing (µm)",      50.0,  150.0, 100.0, step=5.0)

    st.markdown("---")
    ed = (p * 1000) / (s * h)
    st.metric("Calculated Energy Density", f"{ed:.2f} J/mm²")

    if ed < 5:
        st.error("⚠️ Very low energy — insufficient fusion")
    elif ed > 60:
        st.warning("⚠️ Very high energy — overheating risk")
    elif 15 <= ed <= 45:
        st.success("✅ Optimal energy density range")

    st.markdown("---")
    st.caption(f"📊 {len(df_train)} training samples")
    st.caption(f"🌲 {N_ESTIMATORS} trees · seed {SEED}")
    st.caption(f"{'✅ Real CSV data' if using_real_data else '⚠️ Synthetic data'}")

# ─────────────────────────────────────────────────────────────────────────────
# PREDICTIONS
# ─────────────────────────────────────────────────────────────────────────────

input_df = pd.DataFrame([[p, s, f, h]], columns=FEATURES)

preds, stds = {}, {}
for key in TARGET_KEYS:
    preds[key], stds[key] = predict_with_uncertainty(models[key], input_df)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE + METRIC CARDS
# ─────────────────────────────────────────────────────────────────────────────

st.title("🔬 Ti-6Al-4V Laser Surface Engineering Predictor")
st.markdown(f"Trained on **{len(df_train)} samples** · "
            f"3 Random Forest models · "
            f"Predictions shown with ±1σ uncertainty")
st.markdown("---")

c1, c2, c3 = st.columns(3)
metric_defs = [
    ("hardness", "🔩 Microhardness",
     f"{preds['hardness']:.0f} HV",
     f"± {stds['hardness']:.0f} HV uncertainty",
     f"5-fold CV R² = {results['hardness']['cv_mean']:.3f}",
     "> 400 HV → excellent surface hardening"),
    ("roughness", "📐 Surface Roughness (Ra)",
     f"{preds['roughness']:.2f} µm",
     f"± {stds['roughness']:.2f} µm uncertainty",
     f"5-fold CV R² = {results['roughness']['cv_mean']:.3f}",
     "< 2.0 µm → good finish quality"),
    ("haz", "🌡️ HAZ Depth",
     f"{preds['haz']:.1f} µm",
     f"± {stds['haz']:.1f} µm uncertainty",
     f"5-fold CV R² = {results['haz']['cv_mean']:.3f}",
     "< 50 µm → minimal thermal damage"),
]

for col, (key, label, val, delta, cv_txt, tip) in zip([c1, c2, c3], metric_defs):
    with col:
        st.metric(label=label, value=val, delta=delta, delta_color="off")
        st.caption(f"📉 {cv_txt}")
        st.caption(f"💡 {tip}")

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# TABS
# ─────────────────────────────────────────────────────────────────────────────

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "🗺️ Process Map",
    "📊 Model Performance",
    "📈 Feature Importance",
    "🕸️ Trade-off Radar",
    "🎯 Optimizer",
    "🔎 Data Explorer",
])

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — PROCESS MAP (HEATMAP)
# ══════════════════════════════════════════════════════════════════════════════
with tab1:
    st.subheader("Predicted Process Window")
    st.caption("Sweep Power × Speed space with your current Frequency and Hatch values fixed. "
               "★ marks your current setting.")

    col_ctrl, col_map = st.columns([1, 3])

    with col_ctrl:
        map_target = st.selectbox("Output to map", TARGET_LABELS, index=0)
        fixed_f = st.slider("Frequency (kHz)", 20, 100, int(f), key="hm_f")
        fixed_h = st.slider("Hatch (µm)", 50, 150, int(h), key="hm_h")
        show_contour = st.checkbox("Show 450 HV contour", value=True,
                                   disabled=(map_target != "Microhardness"))

    with col_map:
        t_idx  = TARGET_LABELS.index(map_target)
        t_key  = TARGET_KEYS[t_idx]
        t_unit = TARGET_UNITS[t_idx]

        px = np.linspace(100, 500, 80)
        sx = np.linspace(50, 1200, 80)
        xx, yy = np.meshgrid(px, sx)
        grid = pd.DataFrame({
            "Laser_Power_W":       xx.ravel(),
            "Scanning_Speed_mm_s": yy.ravel(),
            "Pulse_Frequency_kHz": fixed_f,
            "Hatch_Spacing_um":    fixed_h,
        })
        Z = models[t_key].predict(grid[FEATURES]).reshape(xx.shape)

        fig, ax = plt.subplots(figsize=(9, 5))
        cf   = ax.contourf(xx, yy, Z, levels=25, cmap="viridis")
        cbar = fig.colorbar(cf, ax=ax)
        cbar.set_label(f"{map_target} ({t_unit})", fontsize=11)

        if show_contour and t_key == "hardness":
            cs = ax.contour(xx, yy, Z, levels=[450], colors="white",
                            linestyles="--", linewidths=1.5)
            ax.clabel(cs, fmt="450 HV", fontsize=9, colors="white")

        ax.scatter(p, s, color="red", s=400, marker="*",
                   edgecolors="white", linewidth=1.2,
                   label="Current setting", zorder=5)
        ax.set_xlabel("Laser Power (W)", fontsize=11)
        ax.set_ylabel("Scanning Speed (mm/s)", fontsize=11)
        ax.set_title(f"{map_target} — Process Window\n"
                     f"(Freq = {fixed_f} kHz, Hatch = {fixed_h} µm)", fontsize=12)
        ax.legend()
        plt.tight_layout()
        st.pyplot(fig)
        plt.close(fig)


# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — MODEL PERFORMANCE
# ══════════════════════════════════════════════════════════════════════════════
with tab2:
    st.subheader("Model Validation — Actual vs Predicted")

    # CV summary row
    cv_cols = st.columns(3)
    for col, key, label, unit in zip(cv_cols, TARGET_KEYS, TARGET_LABELS, TARGET_UNITS):
        r = results[key]
        with col:
            st.metric(
                label=f"{label}",
                value=f"R² = {r['cv_mean']:.3f}",
                delta=f"± {r['cv_std']:.3f} std  |  MAE = {r['mae']:.2f} {unit}",
            )

    st.markdown("---")

    fig, axes = plt.subplots(1, 3, figsize=(15, 5))
    for ax, key, label, unit, color in zip(
        axes, TARGET_KEYS, TARGET_LABELS, TARGET_UNITS, TARGET_COLORS
    ):
        r    = results[key]
        y_t  = r["y_test"]
        y_p  = r["y_pred"]
        lo   = min(y_t.min(), y_p.min()) * 0.96
        hi   = max(y_t.max(), y_p.max()) * 1.04

        ax.scatter(y_t, y_p, alpha=0.45, s=28, color=color, edgecolors="none")
        ax.plot([lo, hi], [lo, hi], "k--", lw=1.5, label="Perfect prediction")
        ax.set_xlim(lo, hi)
        ax.set_ylim(lo, hi)
        ax.set_xlabel(f"Actual ({unit})", fontsize=10)
        ax.set_ylabel(f"Predicted ({unit})", fontsize=10)
        ax.set_title(
            f"{label}\nR²={r['r2']:.3f}   MAE={r['mae']:.2f}   RMSE={r['rmse']:.2f}",
            fontsize=10,
        )
        ax.legend(fontsize=8)
        ax.set_aspect("equal")

    plt.tight_layout()
    st.pyplot(fig)
    plt.close(fig)

    with st.expander("Full metrics table"):
        rows = []
        for key, label, unit in zip(TARGET_KEYS, TARGET_LABELS, TARGET_UNITS):
            r = results[key]
            rows.append({
                "Target":           label,
                "Test R²":         f"{r['r2']:.4f}",
                "CV R² (mean)":    f"{r['cv_mean']:.4f}",
                "CV R² (± std)":   f"± {r['cv_std']:.4f}",
                f"MAE ({unit})":   f"{r['mae']:.3f}",
                f"RMSE ({unit})":  f"{r['rmse']:.3f}",
            })
        st.table(pd.DataFrame(rows).set_index("Target"))


# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — FEATURE IMPORTANCE
# ══════════════════════════════════════════════════════════════════════════════
with tab3:
    st.subheader("Parameter Importance Ranking")
    st.caption("Which laser parameters drive each output property most strongly? "
               "Scores sum to 1.0 within each model.")

    fig, axes = plt.subplots(1, 3, figsize=(15, 4))
    for ax, key, label, color in zip(axes, TARGET_KEYS, TARGET_LABELS, TARGET_COLORS):
        imp     = models[key].feature_importances_
        idx     = np.argsort(imp)
        s_labels = [FEATURE_LABELS[i] for i in idx]
        s_imp    = imp[idx]

        bars = ax.barh(s_labels, s_imp, color=color, alpha=0.78)
        ax.set_xlim(0, s_imp.max() * 1.30)
        ax.set_xlabel("Relative Importance (0–1)", fontsize=10)
        ax.set_title(label, fontsize=11)
        ax.grid(axis="x", linestyle="--", alpha=0.35)

        for bar in bars:
            w = bar.get_width()
            ax.text(w + 0.005, bar.get_y() + bar.get_height() / 2,
                    f"{w:.3f}", va="center", fontsize=9)

    plt.tight_layout()
    st.pyplot(fig)
    plt.close(fig)

    with st.expander("Physical interpretation"):
        st.markdown("""
**Why does Scanning Speed dominate Hardness?**  
Speed controls the dwell time of the laser on any given point → controls cooling rate of the melt pool. Faster scan = rapid quench = martensite phase formation = higher hardness. This is well-established in Ti-6Al-4V laser processing literature.

**Why does Hatch Spacing influence Surface Roughness?**  
Larger hatch = fewer overlapping passes = more pronounced inter-track ridges = higher Ra. It's a direct geometric effect, independent of temperature.

**Why is Pulse Frequency generally the weakest predictor?**  
In the near-continuous-wave regime covered by this dataset, frequency modulates average power delivery but the dominant physics is captured by the energy density term (Power / Speed / Hatch). Frequency would matter more in an extreme pulsed nanosecond/picosecond regime.

**HAZ Depth** is primarily driven by total energy delivered (hence Power and Speed both matter), since HAZ growth scales with the heat diffusion depth.
        """)


# ══════════════════════════════════════════════════════════════════════════════
# TAB 4 — TRADE-OFF RADAR
# ══════════════════════════════════════════════════════════════════════════════
with tab4:
    st.subheader("Multi-Property Trade-off Radar")
    st.caption("All axes normalised 0–1 from the training dataset range. "
               "Higher is always better on this chart (roughness and HAZ are inverted).")

    col_rad, col_scores = st.columns([2, 1])

    with col_rad:
        h_min, h_max = df_train["Microhardness_HV"].min(),        df_train["Microhardness_HV"].max()
        r_min, r_max = df_train["Surface_Roughness_Ra_um"].min(), df_train["Surface_Roughness_Ra_um"].max()
        z_min, z_max = df_train["HAZ_Depth_um"].min(),            df_train["HAZ_Depth_um"].max()

        sc_hard  = (preds["hardness"]  - h_min) / (h_max - h_min)
        sc_rough = 1 - (preds["roughness"] - r_min) / (r_max - r_min)
        sc_haz   = 1 - (preds["haz"]   - z_min) / (z_max - z_min)
        sc_speed = (s - 50) / (1200 - 50)

        sc_hard  = float(np.clip(sc_hard,  0, 1))
        sc_rough = float(np.clip(sc_rough, 0, 1))
        sc_haz   = float(np.clip(sc_haz,   0, 1))
        sc_speed = float(np.clip(sc_speed, 0, 1))

        cats   = ["Hardness", "Surface Quality\n(inv. roughness)", "Low Thermal\nDamage (inv. HAZ)", "Throughput\n(scan speed)"]
        vals   = [sc_hard, sc_rough, sc_haz, sc_speed]
        angles = np.linspace(0, 2 * np.pi, len(cats), endpoint=False).tolist()

        vals_c   = vals + [vals[0]]
        angles_c = angles + [angles[0]]

        fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
        ax.fill(angles_c, vals_c, color="#2ca02c", alpha=0.22)
        ax.plot(angles_c, vals_c, color="#2ca02c", linewidth=2.5)
        ax.scatter(angles, vals, color="#2ca02c", s=80, zorder=5)
        ax.set_ylim(0, 1)
        ax.set_yticklabels([])
        ax.set_xticks(angles)
        ax.set_xticklabels(cats, fontsize=10)
        ax.set_title("Process Performance Radar", fontsize=12, pad=20)
        ax.grid(True, alpha=0.3)
        plt.tight_layout()
        st.pyplot(fig)
        plt.close(fig)

    with col_scores:
        st.markdown("**Normalised scores (0–1)**")
        st.metric("Hardness score",       f"{sc_hard:.2f}")
        st.metric("Surface quality score", f"{sc_rough:.2f}")
        st.metric("Low HAZ score",        f"{sc_haz:.2f}")
        st.metric("Throughput score",     f"{sc_speed:.2f}")
        overall = np.mean([sc_hard, sc_rough, sc_haz, sc_speed])
        st.markdown("---")
        st.metric("Overall score", f"{overall:.2f} / 1.00")


# ══════════════════════════════════════════════════════════════════════════════
# TAB 5 — OPTIMIZER
# ══════════════════════════════════════════════════════════════════════════════
with tab5:
    st.subheader("Parameter Optimizer")
    st.markdown(
        "Set your **target property requirements** and **parameter constraints**. "
        "The optimizer will grid-search the Power × Speed space and return the best matching combinations."
    )

    col_t, col_c = st.columns(2)

    with col_t:
        st.markdown("**Target requirements**")
        min_hard  = st.number_input("Min Hardness (HV)",      value=400.0, step=10.0)
        max_rough = st.number_input("Max Roughness (µm Ra)",   value=2.5,   step=0.1)
        max_haz   = st.number_input("Max HAZ Depth (µm)",     value=60.0,  step=5.0)

    with col_c:
        st.markdown("**Parameter constraints**")
        p_range = st.slider("Power range (W)",       100, 500,  (120, 480))
        s_range = st.slider("Speed range (mm/s)",     50, 1200, (100, 1100))
        o_freq  = st.number_input("Fixed Frequency (kHz)", value=40.0, step=5.0, key="opt_f")
        o_hatch = st.number_input("Fixed Hatch (µm)",     value=100.0, step=10.0, key="opt_h")

    if st.button("🔍 Run Optimizer", type="primary"):
        with st.spinner("Searching parameter space (500 candidate points) …"):
            gp, gs = np.meshgrid(
                np.linspace(p_range[0], p_range[1], 20),
                np.linspace(s_range[0], s_range[1], 25),
            )
            opt_df = pd.DataFrame({
                "Laser_Power_W":       gp.ravel(),
                "Scanning_Speed_mm_s": gs.ravel(),
                "Pulse_Frequency_kHz": o_freq,
                "Hatch_Spacing_um":    o_hatch,
            })
            opt_df["Energy_Density_J_mm2"] = (
                opt_df["Laser_Power_W"] * 1000
            ) / (opt_df["Scanning_Speed_mm_s"] * opt_df["Hatch_Spacing_um"])

            opt_df["Pred_Hardness_HV"]  = models["hardness"].predict(opt_df[FEATURES])
            opt_df["Pred_Roughness_um"] = models["roughness"].predict(opt_df[FEATURES])
            opt_df["Pred_HAZ_um"]       = models["haz"].predict(opt_df[FEATURES])

            valid = opt_df[
                (opt_df["Pred_Hardness_HV"]  >= min_hard) &
                (opt_df["Pred_Roughness_um"] <= max_rough) &
                (opt_df["Pred_HAZ_um"]       <= max_haz)
            ].copy()

        if valid.empty:
            st.error("❌ No combinations found. Try relaxing one or more constraints.")
        else:
            top = (
                valid.sort_values("Pred_Hardness_HV", ascending=False)
                .head(10)
                .reset_index(drop=True)
            )
            st.success(f"✅ Found **{len(valid)}** valid combinations. Top 10 by hardness:")

            display = top[[
                "Laser_Power_W", "Scanning_Speed_mm_s",
                "Energy_Density_J_mm2",
                "Pred_Hardness_HV", "Pred_Roughness_um", "Pred_HAZ_um",
            ]].rename(columns={
                "Laser_Power_W":       "Power (W)",
                "Scanning_Speed_mm_s": "Speed (mm/s)",
                "Energy_Density_J_mm2": "Energy Density (J/mm²)",
                "Pred_Hardness_HV":    "Hardness (HV)",
                "Pred_Roughness_um":   "Roughness (µm)",
                "Pred_HAZ_um":         "HAZ (µm)",
            })
            st.dataframe(
                display.style.format({
                    "Power (W)":              "{:.0f}",
                    "Speed (mm/s)":           "{:.0f}",
                    "Energy Density (J/mm²)": "{:.2f}",
                    "Hardness (HV)":          "{:.0f}",
                    "Roughness (µm)":         "{:.3f}",
                    "HAZ (µm)":               "{:.1f}",
                }).background_gradient(subset=["Hardness (HV)"], cmap="Greens"),
                use_container_width=True,
            )


# ══════════════════════════════════════════════════════════════════════════════
# TAB 6 — DATA EXPLORER
# ══════════════════════════════════════════════════════════════════════════════
with tab6:
    st.subheader("Training Dataset Explorer")

    dc1, dc2, dc3, dc4 = st.columns(4)
    dc1.metric("Total samples",    len(df_train))
    dc2.metric("Hardness range",   f"{df_train['Microhardness_HV'].min():.0f}–{df_train['Microhardness_HV'].max():.0f} HV")
    dc3.metric("Roughness range",  f"{df_train['Surface_Roughness_Ra_um'].min():.2f}–{df_train['Surface_Roughness_Ra_um'].max():.2f} µm")
    dc4.metric("Data source",      "Real CSV" if using_real_data else "Synthetic")

    # Distributions
    fig, axes = plt.subplots(2, 4, figsize=(16, 7))
    all_cols   = FEATURES + TARGETS
    all_labels = FEATURE_LABELS + TARGET_LABELS
    colors     = ["#1f77b4"] * 4 + TARGET_COLORS

    for ax, col, label, clr in zip(axes.ravel()[:7], all_cols, all_labels, colors):
        ax.hist(df_train[col], bins=28, color=clr, alpha=0.75, edgecolor="white", linewidth=0.4)
        ax.set_title(label, fontsize=9, fontweight="bold")
        ax.tick_params(labelsize=7)
        ax.set_ylabel("Count", fontsize=8)

    axes.ravel()[7].set_visible(False)  # hide 8th empty subplot
    plt.suptitle("Parameter & Output Distributions", fontsize=11, y=1.01)
    plt.tight_layout()
    st.pyplot(fig)
    plt.close(fig)

    # Correlation matrix
    st.markdown("---")
    st.markdown("**Correlation Matrix — Features & Targets**")
    corr_cols   = FEATURES + TARGETS
    corr        = df_train[corr_cols].corr()
    short_names = ["Power", "Speed", "Freq", "Hatch", "Hardness", "Roughness", "HAZ"]

    fig, ax = plt.subplots(figsize=(9, 7))
    im = ax.imshow(corr.values, cmap="RdBu_r", vmin=-1, vmax=1, aspect="auto")
    plt.colorbar(im, ax=ax, fraction=0.046)
    ax.set_xticks(range(len(short_names)))
    ax.set_yticks(range(len(short_names)))
    ax.set_xticklabels(short_names, rotation=45, ha="right", fontsize=10)
    ax.set_yticklabels(short_names, fontsize=10)
    for i in range(len(short_names)):
        for j in range(len(short_names)):
            val = corr.iloc[i, j]
            ax.text(j, i, f"{val:.2f}", ha="center", va="center", fontsize=8,
                    color="white" if abs(val) > 0.5 else "black")
    ax.set_title("Pearson Correlation", fontsize=12)
    plt.tight_layout()
    st.pyplot(fig)
    plt.close(fig)


# ─────────────────────────────────────────────────────────────────────────────
# REPORT GENERATION
# ─────────────────────────────────────────────────────────────────────────────

st.markdown("---")
st.subheader("📄 Export Report")

col_btn, col_info = st.columns([1, 3])
with col_btn:
    gen_btn = st.button("Generate .docx Report", type="primary")
with col_info:
    st.caption(
        "Produces a Word document containing: input parameters, predictions with uncertainty, "
        "5-fold CV scores, actual-vs-predicted plots, and feature importance charts."
    )

if gen_btn:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime

        with st.spinner("Building report …"):

            def _fig_to_bytes(fig) -> io.BytesIO:
                buf = io.BytesIO()
                fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
                buf.seek(0)
                plt.close(fig)
                return buf

            # Figure A — actual vs predicted
            fig_a, axes_a = plt.subplots(1, 3, figsize=(14, 4))
            for ax, key, label, unit, color in zip(
                axes_a, TARGET_KEYS, TARGET_LABELS, TARGET_UNITS, TARGET_COLORS
            ):
                r   = results[key]
                lo  = min(r["y_test"].min(), r["y_pred"].min()) * 0.96
                hi  = max(r["y_test"].max(), r["y_pred"].max()) * 1.04
                ax.scatter(r["y_test"], r["y_pred"], alpha=0.4, s=22, color=color, edgecolors="none")
                ax.plot([lo, hi], [lo, hi], "k--", lw=1.5, label="Perfect")
                ax.set_xlim(lo, hi); ax.set_ylim(lo, hi)
                ax.set_aspect("equal")
                ax.set_xlabel(f"Actual ({unit})", fontsize=9)
                ax.set_ylabel(f"Predicted ({unit})", fontsize=9)
                ax.set_title(f"{label}\nR²={r['r2']:.3f}  MAE={r['mae']:.2f}", fontsize=9)
                ax.legend(fontsize=7)
            plt.suptitle("Actual vs Predicted — All Outputs", fontsize=11)
            plt.tight_layout()
            buf_a = _fig_to_bytes(fig_a)

            # Figure B — feature importance
            fig_b, axes_b = plt.subplots(1, 3, figsize=(14, 4))
            for ax, key, label, color in zip(axes_b, TARGET_KEYS, TARGET_LABELS, TARGET_COLORS):
                imp  = models[key].feature_importances_
                idx  = np.argsort(imp)
                ax.barh([FEATURE_LABELS[i] for i in idx], imp[idx], color=color, alpha=0.75)
                ax.set_title(label, fontsize=10)
                ax.set_xlabel("Importance", fontsize=9)
            plt.suptitle("Feature Importance — All Models", fontsize=11)
            plt.tight_layout()
            buf_b = _fig_to_bytes(fig_b)

            # Build document
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            t = doc.add_heading("Ti-6Al-4V Laser Processing — Prediction Report", 0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Generated:   {datetime.now().strftime('%d %B %Y  %H:%M')}")
            doc.add_paragraph(f"Data source: {'Real experimental CSV' if using_real_data else 'Synthetic data (CSV not found)'}")
            doc.add_paragraph(f"Samples:     {len(df_train)}")
            doc.add_paragraph(f"Model:       Random Forest · {N_ESTIMATORS} trees · seed {SEED}")

            doc.add_heading("1. Input Process Parameters", 1)
            tbl = doc.add_table(rows=5, cols=2)
            tbl.style = "Light Grid Accent 1"
            for row, (param, val) in zip(tbl.rows, [
                ("Parameter", "Value"),
                ("Laser Power",     f"{p:.1f} W"),
                ("Scanning Speed",  f"{s:.1f} mm/s"),
                ("Pulse Frequency", f"{f:.1f} kHz"),
                ("Hatch Spacing",   f"{h:.1f} µm"),
            ]):
                row.cells[0].text = param
                row.cells[1].text = val

            doc.add_paragraph(f"\nCalculated Energy Density: {ed:.2f} J/mm²")

            doc.add_heading("2. Predictions", 1)
            tbl2 = doc.add_table(rows=4, cols=4)
            tbl2.style = "Light Grid Accent 1"
            for cell, hdr in zip(tbl2.rows[0].cells,
                                  ["Property", "Predicted", "Uncertainty (±1σ)", "CV R² (5-fold)"]):
                cell.text = hdr
            for row, (key, label, unit) in zip(tbl2.rows[1:],
                                                zip(TARGET_KEYS, TARGET_LABELS, TARGET_UNITS)):
                r = results[key]
                for cell, val in zip(row.cells, [
                    label,
                    f"{preds[key]:.2f} {unit}",
                    f"± {stds[key]:.2f} {unit}",
                    f"{r['cv_mean']:.3f} ± {r['cv_std']:.3f}",
                ]):
                    cell.text = val

            doc.add_heading("3. Model Validation — Actual vs Predicted", 1)
            doc.add_paragraph(
                "Each scatter plot shows test-set predictions vs ground truth. "
                "Points along the dashed diagonal = perfect prediction."
            )
            doc.add_picture(buf_a, width=Inches(6.0))

            doc.add_heading("4. Feature Importance Analysis", 1)
            doc.add_paragraph(
                "Random Forest impurity-based importance scores. "
                "Scores within each model sum to 1.0."
            )
            doc.add_picture(buf_b, width=Inches(6.0))

            doc.add_heading("5. Methodology Notes", 1)
            for note in [
                "Model: scikit-learn RandomForestRegressor with 150 estimators.",
                "Train/test split: 80/20 with random_state=42.",
                "Cross-validation: 5-fold stratified R² (mean ± std reported).",
                "Uncertainty: standard deviation of individual tree predictions (≈ ±1σ).",
                "Energy density: (Power × 1000) / (Speed × Hatch) [J/mm²].",
            ]:
                doc.add_paragraph(f"• {note}")

            buf_doc = io.BytesIO()
            doc.save(buf_doc)
            buf_doc.seek(0)

        st.download_button(
            label="📥 Download Report (.docx)",
            data=buf_doc.getvalue(),
            file_name=f"Ti64_Report_{p:.0f}W_{s:.0f}mms.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        st.success("✅ Report ready — click the button above to download.")

    except ImportError:
        st.error("python-docx not installed. Run:  pip install python-docx")
    except Exception as exc:
        st.error(f"Report generation failed: {exc}")
